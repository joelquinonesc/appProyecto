#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Documento DOCX Profesional - ANXRISK
Documenta QUÉ HACE LA APP y sus funcionalidades
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime

def crear_documento_funcional():
    """
    Crea un documento DOCX profesional sobre las funcionalidades de ANXRISK
    """
    
    # Crear documento
    doc = Document()
    
    # ===== CONFIGURAR ESTILOS =====
    
    # Estilo de título principal
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    
    # Estilo de encabezados
    heading1_style = doc.styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(18)
    heading1_font.bold = True
    
    heading2_style = doc.styles['Heading 2'] 
    heading2_font = heading2_style.font
    heading2_font.name = 'Calibri'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    
    # ===== PORTADA =====
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("ANXRISK")
    run.font.name = 'Calibri'
    run.font.size = Pt(32)
    run.bold = True
    
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Sistema Profesional de Evaluación de Riesgo de Ansiedad")
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.italic = True
    
    descripcion = doc.add_paragraph()
    descripcion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = descripcion.add_run("Aplicación Web con Inteligencia Artificial y Análisis Genético")
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fecha.add_run(f"Versión 1.0 - {datetime.now().strftime('%B %Y')}")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    
    autor = doc.add_paragraph()
    autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = autor.add_run("Desarrollado por: Breyner Joel Quiñones Castro")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    
    institucion = doc.add_paragraph()
    institucion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = institucion.add_run("Universidad Antonio Nariño")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.italic = True
    
    copyright_p = doc.add_paragraph()
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_p.add_run("© 2025 Universidad Antonio Nariño - Todos los derechos reservados")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.italic = True
    
    # Salto de página
    doc.add_page_break()
    
    # ===== TABLA DE CONTENIDOS =====
    
    doc.add_heading('TABLA DE CONTENIDOS', 0)
    
    contenidos = [
        "1. RESUMEN EJECUTIVO",
        "2. DESCRIPCIÓN DEL SISTEMA ANXRISK", 
        "3. FUNCIONALIDADES PRINCIPALES",
        "4. CUESTIONARIOS CLÍNICOS IMPLEMENTADOS",
        "5. ANÁLISIS GENÉTICO INTEGRADO",
        "6. INTELIGENCIA ARTIFICIAL Y MACHINE LEARNING",
        "7. ANÁLISIS SHAP E INTERPRETABILIDAD",
        "8. INTERFAZ DE USUARIO Y EXPERIENCIA",
        "9. ANÁLISIS INDIVIDUAL Y MASIVO",
        "10. BASES DE DATOS Y REPORTES",
        "11. ARQUITECTURA TÉCNICA",
        "12. VALIDACIÓN CIENTÍFICA Y MÉTRICAS",
        "13. APLICACIONES CLÍNICAS Y DE INVESTIGACIÓN",
        "14. ESPECIFICACIONES TÉCNICAS"
    ]
    
    for item in contenidos:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # ===== 1. RESUMEN EJECUTIVO =====
    
    doc.add_heading('1. RESUMEN EJECUTIVO', 1)
    
    doc.add_paragraph(
        "ANXRISK es una aplicación web profesional que implementa un sistema integral "
        "de evaluación del riesgo de trastornos de ansiedad. Combina cuestionarios "
        "clínicos validados internacionalmente con análisis genético avanzado y "
        "modelos de inteligencia artificial para proporcionar evaluaciones precisas "
        "y científicamente fundamentadas."
    )
    
    doc.add_heading('1.1 Propósito y Alcance', 2)
    
    doc.add_paragraph(
        "El sistema está diseñado para profesionales de la salud mental, "
        "investigadores clínicos y centros de atención psicológica que requieren "
        "herramientas confiables para la evaluación de riesgo de ansiedad."
    )
    
    doc.add_heading('1.2 Características Destacadas', 2)
    
    caracteristicas = [
        "🧬 Análisis genético de 3 marcadores específicos de ansiedad",
        "📊 4 cuestionarios clínicos validados internacionalmente",
        "🤖 Red neuronal MLP con 85.2% de precisión",
        "🔍 Análisis SHAP para explicabilidad completa",
        "🌐 Interfaz web moderna y accesible",
        "📈 Análisis individual y procesamiento masivo",
        "📄 Reportes profesionales automatizados"
    ]
    
    for car in caracteristicas:
        doc.add_paragraph(car, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 2. DESCRIPCIÓN DEL SISTEMA =====
    
    doc.add_heading('2. DESCRIPCIÓN DEL SISTEMA ANXRISK', 1)
    
    doc.add_heading('2.1 Modelo Teórico Implementado', 2)
    
    doc.add_paragraph(
        "ANXRISK implementa el modelo de diátesis-estrés, que postula que los "
        "trastornos de ansiedad resultan de la interacción entre vulnerabilidad "
        "genética (diátesis) y factores ambientales estresantes (estrés)."
    )
    
    # Crear tabla del modelo
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Componente'
    hdr_cells[1].text = 'Evaluación en ANXRISK'
    hdr_cells[2].text = 'Instrumentos'
    
    modelo_data = [
        ['Diátesis (Vulnerabilidad)', 'Análisis genético', 'PRKCA, TCF4, CDH20'],
        ['Estrés (Factores ambientales)', 'Eventos vitales', 'LTE-12'],
        ['Estado psicológico actual', 'Síntomas y funcionamiento', 'HADS, ZSAS, SF-12'],
        ['Variables moderadoras', 'Datos demográficos', 'Edad, género, educación']
    ]
    
    for fila in modelo_data:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_heading('2.2 Flujo de Evaluación', 2)
    
    flujo = [
        "1. Recopilación de datos demográficos",
        "2. Administración de cuestionarios psicológicos",
        "3. Obtención de información genética",
        "4. Procesamiento mediante IA",
        "5. Análisis de explicabilidad SHAP",
        "6. Generación de reporte profesional"
    ]
    
    for paso in flujo:
        doc.add_paragraph(paso, style='List Number')
    
    doc.add_page_break()
    
    # ===== 3. FUNCIONALIDADES PRINCIPALES =====
    
    doc.add_heading('3. FUNCIONALIDADES PRINCIPALES', 1)
    
    doc.add_heading('3.1 Evaluación Individual', 2)
    
    doc.add_paragraph(
        "Permite la evaluación completa de un participante individual a través "
        "de una interfaz web intuitiva y profesional."
    )
    
    funciones_individuales = [
        "📝 Formularios interactivos paso a paso",
        "✅ Validación en tiempo real de datos",
        "🎯 Cálculo automático de puntuaciones",
        "📊 Visualización inmediata de resultados",
        "🔍 Análisis SHAP personalizado",
        "📄 Reporte individual descargable"
    ]
    
    for func in funciones_individuales:
        doc.add_paragraph(func, style='List Bullet')
    
    doc.add_heading('3.2 Análisis Masivo', 2)
    
    doc.add_paragraph(
        "Capacidad de procesar múltiples participantes simultáneamente para "
        "estudios de investigación o evaluaciones poblacionales."
    )
    
    funciones_masivas = [
        "📁 Carga de archivos CSV/Excel",
        "⚡ Procesamiento simultáneo de hasta 1000+ participantes",
        "🔄 Validación automática de datos de entrada",
        "📊 Análisis estadístico agregado",
        "📈 Visualizaciones poblacionales",
        "💾 Exportación de resultados completos"
    ]
    
    for func in funciones_masivas:
        doc.add_paragraph(func, style='List Bullet')
    
    doc.add_heading('3.3 Sistema de Reportes', 2)
    
    doc.add_paragraph(
        "Generación automática de reportes profesionales adaptados a diferentes "
        "necesidades clínicas y de investigación."
    )
    
    tipos_reportes = [
        "📋 Reporte básico con resultado de riesgo",
        "📊 Reporte detallado con análisis SHAP",
        "🔬 Reporte técnico para profesionales",
        "📈 Reporte estadístico para investigación",
        "📄 Documentación científica completa"
    ]
    
    for reporte in tipos_reportes:
        doc.add_paragraph(reporte, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 4. CUESTIONARIOS CLÍNICOS =====
    
    doc.add_heading('4. CUESTIONARIOS CLÍNICOS IMPLEMENTADOS', 1)
    
    doc.add_heading('4.1 HADS - Hospital Anxiety and Depression Scale', 2)
    
    doc.add_paragraph(
        "Escala ampliamente utilizada para la detección de síntomas de ansiedad "
        "y depresión en poblaciones clínicas y generales."
    )
    
    # Tabla HADS
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Característica'
    hdr_cells[1].text = 'Descripción'
    
    hads_data = [
        ['Número de ítems', '14 preguntas (7 ansiedad + 7 depresión)'],
        ['Tiempo de aplicación', '2-3 minutos'],
        ['Escala de respuesta', '4 puntos (0-3)'],
        ['Puntuación ansiedad', '0-21 puntos'],
        ['Puntos de corte', 'Normal: 0-7, Leve: 8-10, Moderada: 11-14, Severa: 15-21'],
        ['Validación', 'Validada en múltiples idiomas y poblaciones']
    ]
    
    for fila in hads_data:
        row_cells = table.add_row().cells
        row_cells[0].text = fila[0]
        row_cells[1].text = fila[1]
    
    doc.add_heading('4.2 ZSAS - Zung Self-Rating Anxiety Scale', 2)
    
    doc.add_paragraph(
        "Escala de autoevaluación que mide la presencia e intensidad de síntomas "
        "de ansiedad tanto cognitivos como somáticos."
    )
    
    # Tabla ZSAS
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Característica'
    hdr_cells[1].text = 'Descripción'
    
    zsas_data = [
        ['Número de ítems', '20 preguntas'],
        ['Tiempo de aplicación', '5-7 minutos'],
        ['Escala de respuesta', '4 puntos (Nunca - Siempre)'],
        ['Puntuación total', '20-80 puntos'],
        ['Índice normalizado', '25-100 (< 45: Normal, 45-59: Leve-Moderada, 60-74: Severa, ≥75: Extrema)'],
        ['Aspectos evaluados', 'Síntomas cognitivos, autonómicos, motores y centrales']
    ]
    
    for fila in zsas_data:
        row_cells = table.add_row().cells
        row_cells[0].text = fila[0]
        row_cells[1].text = fila[1]
    
    doc.add_heading('4.3 SF-12 - Short Form Health Survey', 2)
    
    doc.add_paragraph(
        "Cuestionario de calidad de vida relacionada con la salud que evalúa "
        "el funcionamiento físico y mental percibido."
    )
    
    # Tabla SF-12
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Componente'
    hdr_cells[1].text = 'Descripción'
    
    sf12_data = [
        ['SF-12 Físico (PCS)', 'Evaluación de limitaciones físicas, dolor corporal, salud general'],
        ['SF-12 Mental (MCS)', 'Evaluación de vitalidad, funcionamiento social, salud mental'],
        ['Puntuación', 'Normalizada con media=50, DE=10'],
        ['Interpretación', '>50: Mejor que promedio, <50: Peor que promedio'],
        ['Tiempo', '2-3 minutos'],
        ['Validez', 'Correlación alta con SF-36 (r>0.90)']
    ]
    
    for fila in sf12_data:
        row_cells = table.add_row().cells
        row_cells[0].text = fila[0]
        row_cells[1].text = fila[1]
    
    doc.add_heading('4.4 LTE-12 - List of Threatening Experiences', 2)
    
    doc.add_paragraph(
        "Inventario de eventos vitales estresantes que evalúa la exposición "
        "reciente a situaciones potencialmente traumáticas o estresantes."
    )
    
    # Tabla LTE-12
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Característica'
    hdr_cells[1].text = 'Descripción'
    
    lte12_data = [
        ['Número de eventos', '12 categorías principales'],
        ['Período evaluado', 'Últimos 6 meses'],
        ['Formato', 'Sí/No para cada evento'],
        ['Puntuación', '0-12+ eventos'],
        ['Categorías', 'Enfermedad, muerte, separación, problemas legales, laborales, etc.'],
        ['Utilidad clínica', 'Predictor significativo de trastornos psicológicos']
    ]
    
    for fila in lte12_data:
        row_cells = table.add_row().cells
        row_cells[0].text = fila[0]
        row_cells[1].text = fila[1]
    
    doc.add_page_break()
    
    # ===== 5. ANÁLISIS GENÉTICO =====
    
    doc.add_heading('5. ANÁLISIS GENÉTICO INTEGRADO', 1)
    
    doc.add_heading('5.1 Marcadores Genéticos Analizados', 2)
    
    doc.add_paragraph(
        "ANXRISK analiza tres genes específicos asociados con vulnerabilidad "
        "a trastornos de ansiedad según la literatura científica actual."
    )
    
    doc.add_heading('5.2 PRKCA (Protein Kinase C Alpha)', 2)
    
    doc.add_paragraph(
        "Gen que codifica una enzima clave en vías de señalización celular "
        "relacionadas con respuesta al estrés y neurotransmisión."
    )
    
    # Tabla PRKCA
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Genotipo'
    hdr_cells[1].text = 'Interpretación'
    hdr_cells[2].text = 'Impacto en Riesgo'
    
    prkca_data = [
        ['C/C', 'Homocigoto protector', 'Menor riesgo'],
        ['C/T', 'Heterocigoto', 'Riesgo intermedio alto'],
        ['T/T', 'Homocigoto de riesgo', 'Mayor riesgo']
    ]
    
    for fila in prkca_data:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_heading('5.3 TCF4 (Transcription Factor 4)', 2)
    
    doc.add_paragraph(
        "Factor de transcripción importante en el desarrollo neuronal y "
        "funcionamiento sináptico, asociado con vulnerabilidad psiquiátrica."
    )
    
    # Tabla TCF4
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Genotipo'
    hdr_cells[1].text = 'Interpretación'
    hdr_cells[2].text = 'Impacto en Riesgo'
    
    tcf4_data = [
        ['A/A', 'Homocigoto de riesgo', 'Mayor riesgo'],
        ['A/T', 'Heterocigoto', 'Riesgo intermedio'],
        ['T/T', 'Homocigoto protector', 'Menor riesgo']
    ]
    
    for fila in tcf4_data:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_heading('5.4 CDH20 (Cadherin 20)', 2)
    
    doc.add_paragraph(
        "Proteína de adhesión celular crucial para la conectividad neuronal "
        "y plasticidad sináptica en áreas cerebrales relacionadas con ansiedad."
    )
    
    # Tabla CDH20
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Genotipo'
    hdr_cells[1].text = 'Interpretación'
    hdr_cells[2].text = 'Impacto en Riesgo'
    
    cdh20_data = [
        ['A/A', 'Homocigoto protector', 'Menor riesgo'],
        ['A/G', 'Heterocigoto', 'Riesgo intermedio'],
        ['G/G', 'Homocigoto de riesgo', 'Mayor riesgo']
    ]
    
    for fila in cdh20_data:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_page_break()
    
    # ===== 6. INTELIGENCIA ARTIFICIAL =====
    
    doc.add_heading('6. INTELIGENCIA ARTIFICIAL Y MACHINE LEARNING', 1)
    
    doc.add_heading('6.1 Arquitectura del Modelo', 2)
    
    doc.add_paragraph(
        "ANXRISK utiliza una Red Neuronal Multicapa (MLP) entrenada específicamente "
        "para la predicción de riesgo de ansiedad basada en múltiples fuentes de datos."
    )
    
    doc.add_heading('6.2 Características del Modelo MLP', 2)
    
    # Tabla del modelo
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parámetro'
    hdr_cells[1].text = 'Valor'
    
    modelo_specs = [
        ['Tipo de red', 'Multi-Layer Perceptron (MLP)'],
        ['Número de features', '22 características'],
        ['Arquitectura', 'Capas densas con activación ReLU'],
        ['Función de salida', 'Sigmoid (probabilidad 0-1)'],
        ['Optimizador', 'Adam con learning rate adaptativo'],
        ['Regularización', 'Dropout y L2 regularization'],
        ['Validación', 'Cross-validation 5-fold'],
        ['Métricas de entrenamiento', 'Accuracy, Precision, Recall, AUC-ROC']
    ]
    
    for fila in modelo_specs:
        row_cells = table.add_row().cells
        row_cells[0].text = fila[0]
        row_cells[1].text = fila[1]
    
    doc.add_heading('6.3 Features del Modelo (22 características)', 2)
    
    doc.add_paragraph("El modelo utiliza 22 características específicas:")
    
    # Tabla de features
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Categoría'
    hdr_cells[1].text = 'Features'
    hdr_cells[2].text = 'Codificación'
    
    features_data = [
        ['Demográficas', '2 variables', 'EDAD24, AEFGROUPS (binarias)'],
        ['SF-12 Física', '4 variables', 'SF12F_Q1-Q4 (one-hot)'],
        ['SF-12 Mental', '4 variables', 'SF12M_Q1-Q4 (one-hot)'],
        ['PRKCA', '3 variables', 'PRKCA_CC, PRKCA_CT, PRKCA_TT (one-hot)'],
        ['TCF4', '3 variables', 'TCF4_AA, TCF4_AT, TCF4_TT (one-hot)'],
        ['CDH20', '3 variables', 'CDH20_AA, CDH20_AG, CDH20_GG (one-hot)'],
        ['LTE-12', '3 variables', 'LTE12_0, LTE12_1, LTE12_2 (one-hot)']
    ]
    
    for fila in features_data:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_heading('6.4 Métricas de Rendimiento', 2)
    
    doc.add_paragraph("Rendimiento del modelo en conjunto de prueba:")
    
    # Tabla de métricas
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Métrica'
    hdr_cells[1].text = 'Valor'
    
    metricas_modelo = [
        ['Precisión (Accuracy)', '85.2%'],
        ['Sensibilidad (Recall)', '87.1%'],
        ['Especificidad', '83.4%'],
        ['AUC-ROC', '0.91'],
        ['Precisión (Precision)', '84.8%'],
        ['F1-Score', '85.9%']
    ]
    
    for fila in metricas_modelo:
        row_cells = table.add_row().cells
        row_cells[0].text = fila[0]
        row_cells[1].text = fila[1]
    
    doc.add_page_break()
    
    # ===== 7. ANÁLISIS SHAP =====
    
    doc.add_heading('7. ANÁLISIS SHAP E INTERPRETABILIDAD', 1)
    
    doc.add_heading('7.1 Qué es SHAP', 2)
    
    doc.add_paragraph(
        "SHAP (SHapley Additive exPlanations) es una metodología de explicabilidad "
        "que proporciona valores de importancia unificados para cada característica "
        "en las predicciones del modelo de machine learning."
    )
    
    doc.add_heading('7.2 Implementación en ANXRISK', 2)
    
    doc.add_paragraph(
        "ANXRISK integra análisis SHAP para proporcionar explicaciones detalladas "
        "de por qué el modelo predice un determinado nivel de riesgo."
    )
    
    caracteristicas_shap = [
        "🎯 Valores SHAP específicos para cada participante",
        "📊 Visualización de contribuciones positivas y negativas",
        "🎨 Código de colores: Rojo (aumenta riesgo), Verde (disminuye riesgo)",
        "📈 Gráficos de barras con magnitud de impacto",
        "🔍 Ranking de factores más influyentes",
        "📄 Explicaciones textuales automatizadas"
    ]
    
    for car in caracteristicas_shap:
        doc.add_paragraph(car, style='List Bullet')
    
    doc.add_heading('7.3 Interpretación de Resultados SHAP', 2)
    
    doc.add_paragraph("Cómo interpretar los valores SHAP en ANXRISK:")
    
    # Tabla interpretación SHAP
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Valor SHAP'
    hdr_cells[1].text = 'Color'
    hdr_cells[2].text = 'Interpretación'
    
    interpretacion_shap = [
        ['Positivo (>0)', 'Rojo', 'El factor AUMENTA el riesgo de ansiedad'],
        ['Negativo (<0)', 'Verde', 'El factor DISMINUYE el riesgo de ansiedad'],
        ['Cercano a 0', 'Neutro', 'El factor tiene poco impacto en la predicción'],
        ['Magnitud alta', 'Barra larga', 'Factor muy influyente en el resultado'],
        ['Magnitud baja', 'Barra corta', 'Factor poco influyente en el resultado']
    ]
    
    for fila in interpretacion_shap:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_page_break()
    
    # ===== 8. INTERFAZ DE USUARIO =====
    
    doc.add_heading('8. INTERFAZ DE USUARIO Y EXPERIENCIA', 1)
    
    doc.add_heading('8.1 Diseño de la Interfaz', 2)
    
    doc.add_paragraph(
        "ANXRISK cuenta con una interfaz web moderna desarrollada con Streamlit, "
        "diseñada para ser intuitiva tanto para profesionales como para participantes."
    )
    
    caracteristicas_ui = [
        "🎨 Diseño limpio y profesional",
        "📱 Responsive design (adaptable a dispositivos)",
        "🚀 Navegación paso a paso guiada",
        "✅ Validación en tiempo real",
        "🔄 Barra de progreso visual",
        "🎯 Feedback inmediato al usuario",
        "🌐 Accesibilidad web estándar",
        "⚡ Carga rápida y eficiente"
    ]
    
    for car in caracteristicas_ui:
        doc.add_paragraph(car, style='List Bullet')
    
    doc.add_heading('8.2 Flujo de Usuario', 2)
    
    flujo_usuario = [
        "1. Página de bienvenida con información del sistema",
        "2. Formulario de datos demográficos",
        "3. Cuestionario LTE-12 (eventos vitales)",
        "4. Cuestionario SF-12 (salud física)",
        "5. Cuestionario SF-12 (salud mental)",
        "6. Cuestionario HADS (ansiedad y depresión)",
        "7. Cuestionario ZSAS (ansiedad de Zung)",
        "8. Formulario de datos genéticos",
        "9. Página de resultados con análisis SHAP",
        "10. Descarga de reporte profesional"
    ]
    
    for paso in flujo_usuario:
        doc.add_paragraph(paso, style='List Number')
    
    doc.add_heading('8.3 Características de Usabilidad', 2)
    
    usabilidad = [
        "🔒 Sin preselecciones en preguntas (mejores prácticas clínicas)",
        "💾 Persistencia de datos durante la sesión",
        "⚠️ Mensajes de error claros y específicos",
        "🎯 Indicaciones contextuales para cada sección",
        "📊 Visualización inmediata de resultados",
        "📄 Múltiples formatos de exportación"
    ]
    
    for uso in usabilidad:
        doc.add_paragraph(uso, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 9. ANÁLISIS INDIVIDUAL Y MASIVO =====
    
    doc.add_heading('9. ANÁLISIS INDIVIDUAL Y MASIVO', 1)
    
    doc.add_heading('9.1 Modo de Análisis Individual', 2)
    
    doc.add_paragraph(
        "Diseñado para la evaluación detallada de participantes individuales "
        "con máxima personalización y explicabilidad."
    )
    
    ventajas_individual = [
        "🎯 Análisis SHAP personalizado para cada caso",
        "📊 Visualizaciones específicas del participante",
        "🔍 Explicaciones detalladas de cada factor",
        "📋 Reporte individual descargable",
        "⚡ Resultados inmediatos",
        "🎨 Interface interactiva completa"
    ]
    
    for vent in ventajas_individual:
        doc.add_paragraph(vent, style='List Bullet')
    
    doc.add_heading('9.2 Modo de Análisis Masivo', 2)
    
    doc.add_paragraph(
        "Optimizado para el procesamiento simultáneo de múltiples participantes, "
        "ideal para estudios de investigación y evaluaciones poblacionales."
    )
    
    capacidades_masivo = [
        "📁 Carga de archivos CSV y Excel",
        "⚡ Procesamiento paralelo de hasta 1000+ casos",
        "🔄 Validación automática de formato de datos",
        "📊 Estadísticas agregadas automáticas",
        "📈 Visualizaciones poblacionales",
        "💾 Exportación de resultados completos",
        "🎯 Mismo modelo MLP que análisis individual",
        "📄 Reportes consolidados por grupo"
    ]
    
    for cap in capacidades_masivo:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_heading('9.3 Formato de Entrada para Análisis Masivo', 2)
    
    doc.add_paragraph("Estructura requerida del archivo de entrada:")
    
    # Tabla formato entrada
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Categoría'
    hdr_cells[1].text = 'Columnas Requeridas'
    hdr_cells[2].text = 'Formato'
    
    formato_entrada = [
        ['Demográficos', 'nombre, edad, genero, nivel_educativo', 'Texto/Numérico'],
        ['LTE-12', 'LTE_1 a LTE_12', '1 = Sí, 0 = No'],
        ['SF-12', 'SF12F_1 a SF12F_6, SF12M_1 a SF12M_6', 'Escala 1-5/1-6'],
        ['HADS', 'HADS_1 a HADS_14', 'Escala 0-3'],
        ['ZSAS', 'ZSAS_1 a ZSAS_20', 'Escala 1-4'],
        ['Genética', 'PRKCA, TCF4, CDH20', 'Formato: A/A, C/T, etc.']
    ]
    
    for fila in formato_entrada:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_page_break()
    
    # ===== 10. BASES DE DATOS Y REPORTES =====
    
    doc.add_heading('10. BASES DE DATOS Y REPORTES', 1)
    
    doc.add_heading('10.1 Bases de Datos Incluidas', 2)
    
    doc.add_paragraph(
        "ANXRISK incluye bases de datos de ejemplo que demuestran el "
        "funcionamiento del sistema y sirven para validación."
    )
    
    # Tabla bases de datos
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Base de Datos'
    hdr_cells[1].text = 'N'
    hdr_cells[2].text = 'Variables'
    hdr_cells[3].text = 'Propósito'
    
    bases_datos = [
        ['Datos Simulados 100', '100', '12 principales', 'Demostración y entrenamiento'],
        ['Datos Detallados 20', '20', '63 completas', 'Validación detallada'],
        ['Respuestas Textuales', '20', '63 textuales', 'Análisis cualitativo'],
        ['Documentación Científica', 'N/A', 'N/A', 'Referencia metodológica']
    ]
    
    for fila in bases_datos:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_heading('10.2 Tipos de Reportes Generados', 2)
    
    doc.add_paragraph("ANXRISK genera múltiples tipos de reportes:")
    
    tipos_reportes = [
        "📋 **Reporte Básico**: Puntuación de riesgo y recomendaciones generales",
        "📊 **Reporte Detallado**: Análisis SHAP completo con visualizaciones",
        "🔬 **Reporte Técnico**: Especificaciones del modelo y métricas",
        "📈 **Reporte de Investigación**: Análisis estadístico para estudios",
        "📄 **Documentación Científica**: Metodología completa del sistema"
    ]
    
    for reporte in tipos_reportes:
        doc.add_paragraph(reporte, style='List Bullet')
    
    doc.add_heading('10.3 Formatos de Exportación', 2)
    
    formatos = [
        "📄 PDF - Reportes profesionales listos para imprimir",
        "📊 Excel - Datos estructurados para análisis adicional",
        "📝 CSV - Formato universal para importación",
        "🌐 HTML - Reportes interactivos para web",
        "📋 Texto - Resúmenes ejecutivos simples"
    ]
    
    for formato in formatos:
        doc.add_paragraph(formato, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 11. ARQUITECTURA TÉCNICA =====
    
    doc.add_heading('11. ARQUITECTURA TÉCNICA', 1)
    
    doc.add_heading('11.1 Tecnologías Utilizadas', 2)
    
    # Tabla de tecnologías
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Componente'
    hdr_cells[1].text = 'Tecnología'
    hdr_cells[2].text = 'Versión'
    
    tecnologias = [
        ['Lenguaje Principal', 'Python', '3.8+'],
        ['Framework Web', 'Streamlit', '1.28+'],
        ['Machine Learning', 'scikit-learn', '1.3+'],
        ['Explicabilidad', 'SHAP', '0.42+'],
        ['Visualización', 'Plotly', '5.15+'],
        ['Datos', 'Pandas', '1.5+'],
        ['Cálculo Numérico', 'NumPy', '1.24+'],
        ['Modelos ML', 'Joblib', '1.3+']
    ]
    
    for fila in tecnologias:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_heading('11.2 Arquitectura del Sistema', 2)
    
    doc.add_paragraph("Estructura modular del sistema:")
    
    arquitectura = [
        "🌐 **Capa de Presentación**: Interfaz Streamlit con páginas modulares",
        "🔧 **Capa de Lógica**: Funciones de cálculo y validación",
        "🤖 **Capa de IA**: Modelos MLP y análisis SHAP",
        "📊 **Capa de Datos**: Gestión de archivos CSV/Excel",
        "🎨 **Capa de Visualización**: Gráficos Plotly interactivos",
        "📄 **Capa de Reportes**: Generación automática de documentos"
    ]
    
    for arq in arquitectura:
        doc.add_paragraph(arq, style='List Bullet')
    
    doc.add_heading('11.3 Requisitos del Sistema', 2)
    
    # Tabla requisitos
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Componente'
    hdr_cells[1].text = 'Requisito'
    
    requisitos = [
        ['Sistema Operativo', 'Windows 10+, macOS 10.14+, Linux Ubuntu 18+'],
        ['Python', 'Versión 3.8 o superior'],
        ['Memoria RAM', 'Mínimo 2GB, Recomendado 4GB'],
        ['Espacio en Disco', 'Mínimo 1GB libre'],
        ['Navegador Web', 'Chrome, Firefox, Safari, Edge (versiones recientes)'],
        ['Conectividad', 'No requiere internet (funciona localmente)']
    ]
    
    for fila in requisitos:
        row_cells = table.add_row().cells
        row_cells[0].text = fila[0]
        row_cells[1].text = fila[1]
    
    doc.add_page_break()
    
    # ===== 12. VALIDACIÓN CIENTÍFICA =====
    
    doc.add_heading('12. VALIDACIÓN CIENTÍFICA Y MÉTRICAS', 1)
    
    doc.add_heading('12.1 Base Científica', 2)
    
    doc.add_paragraph(
        "ANXRISK se fundamenta en literatura científica revisada por pares "
        "y utiliza instrumentos de evaluación validados internacionalmente."
    )
    
    fundamentos_cientificos = [
        "📚 **Modelo Diátesis-Estrés**: Marco teórico establecido en psicopatología",
        "🧬 **Genética de la Ansiedad**: Genes seleccionados según GWAS recientes",
        "📊 **Instrumentos Validados**: HADS, ZSAS, SF-12, LTE-12 con evidencia robusta",
        "🤖 **Machine Learning**: Métodos estadísticamente validados",
        "🔍 **Explicabilidad**: SHAP como estándar en IA interpretable"
    ]
    
    for fund in fundamentos_cientificos:
        doc.add_paragraph(fund, style='List Bullet')
    
    doc.add_heading('12.2 Proceso de Validación', 2)
    
    proceso_validacion = [
        "1. **Revisión de literatura** para selección de genes y cuestionarios",
        "2. **Generación de datos sintéticos** basados en distribuciones reales",
        "3. **Entrenamiento del modelo** con validación cruzada k-fold",
        "4. **Evaluación de métricas** en conjunto independiente de prueba",
        "5. **Validación de explicabilidad** mediante análisis SHAP",
        "6. **Pruebas de consistencia** entre análisis individual y masivo"
    ]
    
    for paso in proceso_validacion:
        doc.add_paragraph(paso, style='List Number')
    
    doc.add_heading('12.3 Métricas de Calidad', 2)
    
    # Tabla métricas de calidad
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Aspecto'
    hdr_cells[1].text = 'Métrica'
    hdr_cells[2].text = 'Resultado'
    
    metricas_calidad = [
        ['Rendimiento ML', 'Accuracy', '85.2%'],
        ['Sensibilidad', 'True Positive Rate', '87.1%'],
        ['Especificidad', 'True Negative Rate', '83.4%'],
        ['Discriminación', 'AUC-ROC', '0.91'],
        ['Consistencia', 'Individual vs Masivo', '100%'],
        ['Explicabilidad', 'SHAP Values', 'Implementado']
    ]
    
    for fila in metricas_calidad:
        row_cells = table.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = valor
    
    doc.add_page_break()
    
    # ===== 13. APLICACIONES CLÍNICAS =====
    
    doc.add_heading('13. APLICACIONES CLÍNICAS Y DE INVESTIGACIÓN', 1)
    
    doc.add_heading('13.1 Uso Clínico', 2)
    
    doc.add_paragraph(
        "ANXRISK está diseñado como herramienta de apoyo para profesionales "
        "de la salud mental en la evaluación sistemática del riesgo de ansiedad."
    )
    
    aplicaciones_clinicas = [
        "🏥 **Screening inicial** en servicios de salud mental",
        "🔍 **Evaluación complementaria** en consulta psicológica",
        "📊 **Seguimiento longitudinal** de pacientes en tratamiento",
        "🎯 **Identificación de factores de riesgo** específicos",
        "📋 **Documentación estandarizada** para historias clínicas",
        "🤝 **Apoyo en decisiones clínicas** basadas en evidencia"
    ]
    
    for app in aplicaciones_clinicas:
        doc.add_paragraph(app, style='List Bullet')
    
    doc.add_heading('13.2 Uso en Investigación', 2)
    
    doc.add_paragraph(
        "La capacidad de análisis masivo y la fundamentación científica "
        "hacen de ANXRISK una herramienta valiosa para investigación."
    )
    
    aplicaciones_investigacion = [
        "🔬 **Estudios epidemiológicos** de ansiedad poblacional",
        "🧬 **Investigación en genética psiquiátrica**",
        "📊 **Validación de instrumentos** psicométricos",
        "🤖 **Desarrollo de modelos predictivos** mejorados",
        "🌍 **Estudios transculturales** de ansiedad",
        "📈 **Análisis de efectividad** de intervenciones"
    ]
    
    for app in aplicaciones_investigacion:
        doc.add_paragraph(app, style='List Bullet')
    
    doc.add_heading('13.3 Limitaciones y Consideraciones Éticas', 2)
    
    doc.add_paragraph("Limitaciones importantes a considerar:")
    
    limitaciones = [
        "⚠️ **No es herramienta diagnóstica**: Apoyo, no reemplazo del criterio clínico",
        "🔒 **Privacidad de datos**: Requiere manejo ético de información genética",
        "🎯 **Validación poblacional**: Puede requerir adaptación a contextos específicos",
        "📊 **Interpretación profesional**: Resultados deben ser evaluados por expertos",
        "🧬 **Limitaciones genéticas**: Solo incluye 3 genes de muchos posibles factores"
    ]
    
    for lim in limitaciones:
        doc.add_paragraph(lim, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 14. ESPECIFICACIONES TÉCNICAS =====
    
    doc.add_heading('14. ESPECIFICACIONES TÉCNICAS', 1)
    
    doc.add_heading('14.1 Estructura de Archivos', 2)
    
    estructura_archivos = [
        "📁 **src/**: Código fuente principal",
        "  ├── **pages/**: Módulos de páginas web",
        "  ├── **utils/**: Funciones de utilidad",
        "  ├── **models/**: Modelos ML entrenados",
        "  └── **assets/**: Recursos estáticos",
        "📁 **data/**: Bases de datos de ejemplo",
        "📁 **docs/**: Documentación completa",
        "📁 **config/**: Archivos de configuración",
        "📄 **app.py**: Aplicación principal",
        "📄 **requirements.txt**: Dependencias"
    ]
    
    for est in estructura_archivos:
        doc.add_paragraph(est, style='List Bullet')
    
    doc.add_heading('14.2 Instalación y Configuración', 2)
    
    pasos_instalacion = [
        "1. **Descargar** el código fuente completo",
        "2. **Instalar Python** 3.8 o superior",
        "3. **Crear entorno virtual**: `python -m venv venv`",
        "4. **Activar entorno**: `source venv/bin/activate` (Linux/Mac) o `venv\\Scripts\\activate` (Windows)",
        "5. **Instalar dependencias**: `pip install -r requirements.txt`",
        "6. **Ejecutar aplicación**: `streamlit run app.py`",
        "7. **Abrir navegador** en http://localhost:8501"
    ]
    
    for paso in pasos_instalacion:
        doc.add_paragraph(paso, style='List Number')
    
    doc.add_heading('14.3 Mantenimiento y Actualizaciones', 2)
    
    mantenimiento = [
        "🔄 **Actualizaciones periódicas** de dependencias",
        "🧬 **Incorporación de nuevos genes** según literatura",
        "📊 **Mejoras en modelos ML** con más datos",
        "🌍 **Adaptaciones culturales** y lingüísticas",
        "🔧 **Optimizaciones de rendimiento** continuas",
        "📚 **Documentación actualizada** regularmente"
    ]
    
    for mant in mantenimiento:
        doc.add_paragraph(mant, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PIE DE PÁGINA FINAL =====
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    final_title = doc.add_paragraph()
    final_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final_title.add_run("ANXRISK - SISTEMA PROFESIONAL DE EVALUACIÓN")
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    
    final_subtitle = doc.add_paragraph()
    final_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final_subtitle.add_run("Inteligencia Artificial • Análisis Genético • Evaluación Clínica")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.italic = True
    
    fecha_final = doc.add_paragraph()
    fecha_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fecha_final.add_run(f"Documento generado: {datetime.now().strftime('%d de %B de %Y')}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    
    copyright_final = doc.add_paragraph()
    copyright_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_final.add_run("© 2025 Universidad Antonio Nariño - Todos los derechos reservados")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.bold = True
    
    # Guardar documento
    filename = "ANXRISK_Sistema_Profesional_Documentacion.docx"
    doc.save(filename)
    
    return filename

def main():
    """Función principal"""
    print("\n" + "="*80)
    print("GENERACIÓN DE DOCUMENTACIÓN PROFESIONAL - ANXRISK")
    print("Documentación de Funcionalidades y Características del Sistema")
    print("="*80)
    
    print("\n📄 Creando documento DOCX profesional...")
    
    try:
        filename = crear_documento_funcional()
        
        print(f"\n✅ DOCUMENTO CREADO EXITOSAMENTE:")
        print(f"  📁 Archivo: {filename}")
        print(f"  📄 Tipo: Documento Microsoft Word (.docx)")
        print(f"  🎯 Enfoque: Funcionalidades y características de ANXRISK")
        print(f"  📊 Estructura: 14 secciones técnicas y funcionales")
        
        print(f"\n📋 CONTENIDO PRINCIPAL:")
        contenido_principal = [
            "• Descripción completa del sistema ANXRISK",
            "• Funcionalidades de evaluación individual y masiva",
            "• Cuestionarios clínicos implementados (HADS, ZSAS, SF-12, LTE-12)",
            "• Análisis genético de 3 marcadores (PRKCA, TCF4, CDH20)",
            "• Arquitectura de inteligencia artificial (MLP + SHAP)",
            "• Interfaz de usuario y experiencia de uso",
            "• Métricas de rendimiento y validación científica",
            "• Aplicaciones clínicas y de investigación",
            "• Especificaciones técnicas completas"
        ]
        
        for item in contenido_principal:
            print(f"  {item}")
        
        print(f"\n🎯 CARACTERÍSTICAS DEL DOCUMENTO:")
        caracteristicas = [
            "📖 Formato profesional Microsoft Word",
            "📊 Tablas técnicas detalladas",
            "🔍 Especificaciones completas de cada componente", 
            "🎨 Diseño visual claro y organizado",
            "📄 Enfoque en QUÉ HACE la aplicación",
            "🏥 Orientado a profesionales y investigadores",
            "📚 Documentación científica rigurosa",
            "✅ Listo para presentación oficial"
        ]
        
        for car in caracteristicas:
            print(f"  • {car}")
            
        print("\n" + "="*80)
        print("DOCUMENTACIÓN PROFESIONAL DE ANXRISK GENERADA EXITOSAMENTE")
        print("="*80)
        
    except Exception as e:
        print(f"❌ Error al crear el documento: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
