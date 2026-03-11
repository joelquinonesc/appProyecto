"""
Generador del Manual de Usuario ANXRISK — Versión 5
====================================================
Genera un documento DOCX profesional y descriptivo con toda la
información necesaria para que el profesional de salud mental
utilice correctamente la aplicación ANXRISK.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from datetime import datetime
import os

# ── Colores corporativos ──
AZUL_PRIMARIO = RGBColor(0x2B, 0x87, 0xD1)
GRIS_OSCURO = RGBColor(0x2E, 0x2E, 0x2E)
GRIS_MEDIO = RGBColor(0x66, 0x66, 0x66)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
VERDE = RGBColor(0x4C, 0xAF, 0x50)
ROJO = RGBColor(0xE5, 0x39, 0x35)
NARANJA = RGBColor(0xFF, 0x98, 0x00)


def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda de tabla."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = AZUL_PRIMARIO
    return h


def add_body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GRIS_OSCURO
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = GRIS_OSCURO
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_OSCURO
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = BLANCO
        set_cell_shading(cell, "2B87D1")
    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = GRIS_OSCURO
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F7FF")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def generar_manual():
    doc = Document()

    # ── Configuración de márgenes ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("ANXRISK")
    run.font.size = Pt(42)
    run.font.color.rgb = AZUL_PRIMARIO
    run.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Sistema de Estratificación del Riesgo\nde Trastornos de Ansiedad")
    run.font.size = Pt(18)
    run.font.color.rgb = GRIS_MEDIO

    doc.add_paragraph()

    p_manual = doc.add_paragraph()
    p_manual.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_manual.add_run("MANUAL DE USUARIO")
    run.font.size = Pt(24)
    run.font.color.rgb = GRIS_OSCURO
    run.bold = True

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_ver.add_run("Versión 5.0")
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_PRIMARIO

    for _ in range(4):
        doc.add_paragraph()

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_fecha.add_run(f"Fecha de publicación: {datetime.now().strftime('%d de %B de %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_MEDIO

    p_autor = doc.add_paragraph()
    p_autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_autor.add_run("© 2025–2026 Breyner Joel Quiñones Castro")
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_MEDIO

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLA DE CONTENIDOS (manual)
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Tabla de Contenidos", level=1)
    toc_items = [
        "1. Introducción y Propósito",
        "2. Requisitos del Sistema",
        "3. Instalación y Puesta en Marcha",
        "4. Estructura de la Aplicación",
        "5. Flujo de Evaluación Individual",
        "   5.1 Datos Demográficos",
        "   5.2 Eventos Vitales Estresantes (LTE-12)",
        "   5.3 SF-12 Componente Física (PCS)",
        "   5.4 SF-12 Componente Mental (MCS)",
        "   5.5 Escala HADS de Ansiedad",
        "   5.6 Escala de Ansiedad de Zung (ZSAS)",
        "   5.7 Panel Genético (Módulo de Confirmación)",
        "   5.8 Resultados y Reporte PDF",
        "6. Análisis Masivo (CSV)",
        "7. Modelos de Machine Learning",
        "   7.1 Modelo Estándar (13 features)",
        "   7.2 Modelo Extendido (22 features)",
        "8. Interpretabilidad SHAP",
        "9. Clasificación del Nivel de Riesgo",
        "10. Generación del Reporte PDF",
        "11. Dependencias y Versiones",
        "12. Privacidad y Protección de Datos",
        "13. Preguntas Frecuentes",
        "14. Historial de Cambios (Versión 5)",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = GRIS_OSCURO

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. INTRODUCCIÓN
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Introducción y Propósito", level=1)
    add_body(doc, (
        "ANXRISK es un sistema profesional de estratificación del riesgo de trastornos de ansiedad "
        "desarrollado como herramienta de apoyo a la decisión clínica. Combina instrumentos psicométricos "
        "validados internacionalmente con modelos de aprendizaje automático supervisado (XGBoost) para "
        "proporcionar una evaluación multimodal con interpretabilidad individual."
    ))
    add_body(doc, (
        "La aplicación está construida sobre Streamlit y se despliega en Streamlit Community Cloud. "
        "Los datos del paciente se procesan exclusivamente en la sesión del navegador y no se almacenan "
        "en servidores externos, cumpliendo con la Ley 1581 de 2012 de Protección de Datos Personales "
        "de Colombia."
    ))
    add_body(doc, "La evaluación integra las siguientes fuentes de datos:", bold=True)
    add_bullet(doc, "Edad, género, años de educación formal", bold_prefix="Datos demográficos: ")
    add_bullet(doc, "12 eventos estresantes recientes (escala LTE-12, Brugha et al., 1985)", bold_prefix="Datos psicosociales: ")
    add_bullet(doc, "Calidad de vida SF-12 — componentes física y mental (Ware et al., 1996)", bold_prefix="Salud percibida: ")
    add_bullet(doc, "HADS — 7 ítems de ansiedad (Zigmond & Snaith, 1983)", bold_prefix="Ansiedad clínica: ")
    add_bullet(doc, "ZSAS — 20 ítems de Zung (1971)", bold_prefix="Ansiedad somática: ")
    add_bullet(doc, "SNPs PRKCA, TCF4, CDH20 — módulo de confirmación (opcional)", bold_prefix="Panel genético: ")

    add_body(doc, (
        "Importante: ANXRISK es una herramienta de apoyo clínico. Los resultados deben ser interpretados "
        "por un profesional de salud mental dentro del contexto clínico completo del paciente. "
        "No sustituye el juicio clínico ni constituye un diagnóstico."
    ), italic=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. REQUISITOS DEL SISTEMA
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Requisitos del Sistema", level=1)

    add_heading_styled(doc, "2.1 Para uso en línea (Streamlit Cloud)", level=2)
    add_bullet(doc, "Navegador web moderno (Chrome, Firefox, Edge, Safari)")
    add_bullet(doc, "Conexión a internet estable")
    add_bullet(doc, "No requiere instalación de software adicional")

    add_heading_styled(doc, "2.2 Para ejecución local", level=2)
    add_bullet(doc, "Python 3.12.12 (versión recomendada)")
    add_bullet(doc, "pip (gestor de paquetes de Python)")
    add_bullet(doc, "4 GB de RAM mínimo (8 GB recomendado)")
    add_bullet(doc, "500 MB de espacio en disco")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. INSTALACIÓN
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Instalación y Puesta en Marcha", level=1)

    add_heading_styled(doc, "3.1 Ejecución en Streamlit Cloud (recomendado)", level=2)
    add_body(doc, (
        "ANXRISK se despliega automáticamente desde el repositorio de GitHub en Streamlit Community Cloud. "
        "El archivo runtime.txt especifica la versión de Python (3.12.12) y requirements.txt las dependencias exactas. "
        "No se requiere ninguna acción del usuario final más allá de acceder a la URL proporcionada."
    ))

    add_heading_styled(doc, "3.2 Instalación local", level=2)
    add_body(doc, "Paso 1: Clonar el repositorio:")
    add_body(doc, "  git clone https://github.com/joelquinonesc/appProyecto.git", size=10)
    add_body(doc, "  cd appProyecto", size=10)
    add_body(doc, "Paso 2: Crear entorno virtual e instalar dependencias:")
    add_body(doc, "  python -m venv .venv", size=10)
    add_body(doc, "  source .venv/bin/activate    # Linux/Mac", size=10)
    add_body(doc, "  .venv\\Scripts\\activate       # Windows", size=10)
    add_body(doc, "  pip install -r requirements.txt", size=10)
    add_body(doc, "Paso 3: Ejecutar la aplicación:")
    add_body(doc, "  streamlit run app.py", size=10)
    add_body(doc, "La aplicación se abrirá automáticamente en http://localhost:8501")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. ESTRUCTURA DE LA APLICACIÓN
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Estructura de la Aplicación", level=1)
    add_body(doc, "La aplicación se organiza en los siguientes directorios y archivos principales:")

    add_table(doc,
        ["Archivo / Directorio", "Descripción"],
        [
            ["app.py", "Punto de entrada principal. Enruta entre páginas según el estado de la sesión."],
            ["requirements.txt", "Dependencias de Python con versiones fijadas al entorno de entrenamiento."],
            ["runtime.txt", "Versión de Python para Streamlit Cloud (3.12.12)."],
            ["src/config.py", "Configuración central: rutas de modelos, listas de features, opciones de genotipos."],
            ["src/pages/", "Módulos de cada sección del flujo: home, demograficos, eventos_vitales, sf12_fisica, sf12_mental, hads, zsas, datos_geneticos, resultados, analisis_masivo."],
            ["src/utils/calculos.py", "Funciones puras de transformación: edad→grupo, educación→binaria, SF-12→cuartiles, niveles HADS/ZSAS, clasificación de riesgo."],
            ["src/utils/dataframe_manager.py", "Gestor del DataFrame de sesión: almacena, actualiza y exporta datos del paciente."],
            ["src/models/", "Modelos serializados: anxrisk_best_standard.joblib (XGBoost, 13 features) y anxrisk_best_extended.joblib (XGBoost, 22 features)."],
            ["src/assets/styles/main.css", "Hoja de estilos CSS para la interfaz."],
            ["data/", "Bases de datos simuladas para pruebas y análisis masivo."],
            ["docs/", "Documentación técnica y manuales."],
            ["scripts/", "Scripts auxiliares para generación de datos de prueba y análisis SHAP masivo."],
        ],
        col_widths=[5, 12]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. FLUJO DE EVALUACIÓN INDIVIDUAL
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Flujo de Evaluación Individual", level=1)
    add_body(doc, (
        "La evaluación individual sigue un flujo secuencial obligatorio. Cada sección debe completarse "
        "antes de avanzar a la siguiente. El sistema valida automáticamente que todas las respuestas "
        "estén completas antes de permitir la navegación."
    ))

    add_table(doc,
        ["Paso", "Sección", "Instrumento", "Descripción"],
        [
            ["1", "Datos Demográficos", "—", "Nombre, documento, edad, género, años de educación + datos del profesional evaluador"],
            ["2", "Eventos Vitales", "LTE-12", "12 preguntas Sí/No sobre eventos estresantes recientes"],
            ["3", "Salud Física", "SF-12 PCS", "6 ítems sobre percepción de salud física y limitaciones"],
            ["4", "Salud Mental", "SF-12 MCS", "6 ítems sobre bienestar emocional y vitalidad"],
            ["5", "Ansiedad (HADS)", "HADS-A", "7 ítems sobre síntomas de ansiedad en la última semana"],
            ["6", "Ansiedad (ZSAS)", "ZSAS", "20 ítems afectivos y somáticos de ansiedad"],
            ["7", "Panel Genético", "SNPs", "Genotipos PRKCA, TCF4, CDH20 (módulo de confirmación)"],
            ["8", "Resultados", "XGBoost/NB + SHAP", "Predicción, interpretabilidad y descarga de PDF"],
        ],
        col_widths=[1.2, 3.5, 3, 9]
    )

    doc.add_paragraph()

    # 5.1 Demográficos
    add_heading_styled(doc, "5.1 Datos Demográficos", level=2)
    add_body(doc, (
        "Esta sección recopila la información base del paciente necesaria para la estratificación. "
        "Todos los campos marcados con asterisco (*) son obligatorios."
    ))
    add_body(doc, "Campos del paciente:", bold=True)
    add_bullet(doc, "Texto libre. Se incluye en la portada del reporte PDF.", bold_prefix="Nombre completo *: ")
    add_bullet(doc, "Número de identificación del paciente.", bold_prefix="Documento de identidad *: ")
    add_bullet(doc, "Valor numérico entre 1 y 120 años. Se transforma internamente: ≤24 = grupo 0, >24 = grupo 1.", bold_prefix="Edad *: ")
    add_bullet(doc, "Masculino o Femenino. Se codifica como 0 (Masculino) o 1 (Femenino).", bold_prefix="Género *: ")
    add_bullet(doc, "Valor numérico. Máximo permitido = edad − 5 años. Se transforma: ≤14 = 0, ≥15 = 1.", bold_prefix="Años de educación formal *: ")

    add_body(doc, "Datos del profesional evaluador (opcionales pero recomendados):", bold=True)
    add_bullet(doc, "Nombre del profesional que realiza la evaluación.")
    add_bullet(doc, "Cargo o especialidad (Psiquiatra, Psicólogo clínico, etc.).")
    add_bullet(doc, "Institución donde se realiza la evaluación.")
    add_bullet(doc, "Número de tarjeta profesional.")
    add_body(doc, (
        "Estos datos aparecen en el reporte PDF final, en la sección de firma del profesional."
    ), italic=True)

    # 5.2 LTE-12
    add_heading_styled(doc, "5.2 Eventos Vitales Estresantes (LTE-12)", level=2)
    add_body(doc, (
        "El cuestionario LTE-12 (List of Threatening Experiences) evalúa 12 eventos vitales estresantes "
        "que el paciente pueda haber experimentado recientemente. Cada evento se responde con Sí (1) o No (0). "
        "La suma total se clasifica en tres niveles:"
    ))
    add_table(doc,
        ["Puntaje Total", "Clasificación", "Variable del Modelo"],
        [
            ["0 eventos", "Sin eventos estresantes", "LTE12_0 = 1"],
            ["1 evento", "Un evento estresante", "LTE12_1 = 1"],
            ["2 o más eventos", "Múltiples eventos estresantes", "LTE12_2 = 1"],
        ],
        col_widths=[4, 5.5, 5.5]
    )
    add_body(doc, "Referencia: Brugha, T., Bebbington, P., Tennant, C., & Hurry, J. (1985). Psychological Medicine, 15(1), 189-194.", italic=True, size=10)

    # 5.3 SF-12 Física
    add_heading_styled(doc, "5.3 SF-12 — Componente Física (PCS)", level=2)
    add_body(doc, (
        "Evalúa la percepción de salud física del paciente mediante 6 ítems que cubren: estado general "
        "de salud, limitaciones en actividades moderadas, limitaciones al subir escaleras, reducción de "
        "actividades por salud física, incumplimiento de tareas y dolor. El puntaje se clasifica en cuartiles:"
    ))
    add_table(doc,
        ["Puntaje PCS", "Cuartil", "Etiqueta", "Variable del Modelo"],
        [
            ["≤ 15", "Q1", "Salud Física Muy Baja", "SF12F_Q1 = 1"],
            ["16 – 17", "Q2", "Salud Física Baja", "SF12F_Q2 = 1"],
            ["18 – 19", "Q3", "Salud Física Moderada", "SF12F_Q3 = 1"],
            ["≥ 20", "Q4", "Salud Física Excelente", "SF12F_Q4 = 1"],
        ],
        col_widths=[3, 2, 5, 5]
    )

    # 5.4 SF-12 Mental
    add_heading_styled(doc, "5.4 SF-12 — Componente Mental (MCS)", level=2)
    add_body(doc, (
        "Evalúa el bienestar emocional y mental mediante 6 ítems que cubren: limitaciones por problemas "
        "emocionales, cuidado en actividades por estado emocional, dificultades sociales, tranquilidad, "
        "energía y tristeza. El puntaje se clasifica en cuartiles:"
    ))
    add_table(doc,
        ["Puntaje MCS", "Cuartil", "Etiqueta", "Variable del Modelo"],
        [
            ["≤ 15", "Q1", "Salud Mental Muy Baja", "SF12M_Q1 = 1"],
            ["16 – 18", "Q2", "Salud Mental Baja", "SF12M_Q2 = 1"],
            ["19 – 21", "Q3", "Salud Mental Moderada", "SF12M_Q3 = 1"],
            ["≥ 22", "Q4", "Salud Mental Excelente", "SF12M_Q4 = 1"],
        ],
        col_widths=[3, 2, 5, 5]
    )
    add_body(doc, "Referencia: Ware, J. E., Kosinski, M., & Keller, S. D. (1996). Medical Care, 34(3), 220-233.", italic=True, size=10)

    doc.add_page_break()

    # 5.5 HADS
    add_heading_styled(doc, "5.5 Escala HADS de Ansiedad", level=2)
    add_body(doc, (
        "La escala HADS (Hospital Anxiety and Depression Scale) — subescala de ansiedad — consta de "
        "7 ítems que evalúan manifestaciones emocionales y psicológicas de la ansiedad durante la última "
        "semana. Cada ítem se puntúa de 0 a 3. El puntaje total se clasifica:"
    ))
    add_table(doc,
        ["Puntaje HADS", "Clasificación", "Indicador"],
        [
            ["0 – 7", "✅ Riesgo Bajo", "Sin intervención inmediata requerida"],
            ["≥ 8", "⚠️ Riesgo de Ansiedad", "Requiere evaluación profesional"],
        ],
        col_widths=[3, 5, 7]
    )
    add_body(doc, "Referencia: Zigmond, A. S., & Snaith, R. P. (1983). Acta Psychiatrica Scandinavica, 67(6), 361-370.", italic=True, size=10)

    # 5.6 ZSAS
    add_heading_styled(doc, "5.6 Escala de Ansiedad de Zung (ZSAS)", level=2)
    add_body(doc, (
        "La ZSAS consta de 20 ítems que evalúan aspectos afectivos y somáticos de la ansiedad. "
        "5 ítems tienen puntuación invertida (ítems 5, 9, 13, 17, 19). El puntaje bruto se multiplica "
        "por 1.25 para obtener el puntaje normalizado. La clasificación:"
    ))
    add_table(doc,
        ["Puntaje Normalizado", "Clasificación", "Indicador"],
        [
            ["< 36", "✅ Riesgo Bajo", "Sin intervención inmediata requerida"],
            ["≥ 36", "⚠️ Riesgo de Ansiedad", "Requiere evaluación profesional"],
        ],
        col_widths=[4, 5, 7]
    )
    add_body(doc, "Referencia: Zung, W. W. (1971). Psychosomatics, 12(6), 371-379.", italic=True, size=10)

    # 5.7 Panel Genético
    add_heading_styled(doc, "5.7 Panel Genético (Módulo de Confirmación)", level=2)
    add_body(doc, (
        "El panel genético es un módulo de confirmación que permite incorporar información genómica "
        "a la evaluación. Se activa después de completar todos los cuestionarios clínicos (HADS y ZSAS). "
        "La inclusión de datos genéticos hace que el sistema utilice el modelo extendido (22 features) "
        "en lugar del modelo estándar (13 features)."
    ))
    add_body(doc, "Genes evaluados:", bold=True)
    add_table(doc,
        ["Gen", "Nombre Completo", "Función", "Genotipos Posibles"],
        [
            ["PRKCA", "Proteína Quinasa C Alfa", "Regulación del estrés y respuesta emocional", "T/T, C/T, C/C"],
            ["TCF4", "Factor de Transcripción 4", "Desarrollo neuronal y predisposición psiquiátrica", "A/A, A/T, T/T"],
            ["CDH20", "Cadherina 20", "Conectividad neuronal y neurotransmisión", "G/G, G/A, A/A"],
        ],
        col_widths=[2, 4.5, 5.5, 4]
    )
    add_body(doc, (
        "Cada genotipo se codifica mediante one-hot encoding para el modelo. Por ejemplo, si PRKCA = C/T, "
        "entonces PRKCA_C/C = 0, PRKCA_C/T = 1, PRKCA_T/T = 0."
    ))

    # 5.8 Resultados
    add_heading_styled(doc, "5.8 Resultados y Reporte PDF", level=2)
    add_body(doc, (
        "La página de resultados presenta un resumen completo de la evaluación. Tras completar los "
        "cuestionarios clínicos (o también el panel genético), el profesional presiona 'Generar Evaluación' "
        "para ejecutar la predicción del modelo y el análisis SHAP."
    ))
    add_body(doc, "La página muestra:", bold=True)
    add_bullet(doc, "Resumen de todos los cuestionarios completados con puntajes y clasificaciones.")
    add_bullet(doc, "Opción de incluir o no datos genéticos (si no se incluyen, se usa el modelo estándar).")
    add_bullet(doc, "Probabilidad de riesgo con clasificación en tres niveles (Bajo, Moderado, Alto).")
    add_bullet(doc, "Gráfico SHAP waterfall con la contribución individual de cada variable.")
    add_bullet(doc, "Botón de descarga del reporte en formato PDF.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. ANÁLISIS MASIVO
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Análisis Masivo (CSV)", level=1)
    add_body(doc, (
        "ANXRISK permite evaluar múltiples pacientes simultáneamente mediante la carga de un archivo CSV. "
        "Esta funcionalidad es útil para investigadores o instituciones que necesiten procesar lotes de datos."
    ))
    add_body(doc, "El archivo CSV debe contener las columnas requeridas por el modelo seleccionado:", bold=True)
    add_body(doc, "Modelo estándar (13 columnas): EDAD24, AEFGROUPS, LTE12_0, LTE12_1, LTE12_2, SF12F_Q1, SF12F_Q2, SF12F_Q3, SF12F_Q4, SF12M_Q1, SF12M_Q2, SF12M_Q3, SF12M_Q4", size=10)
    add_body(doc, "Modelo extendido (22 columnas): las 13 anteriores + PRKCA_C/C, PRKCA_C/T, PRKCA_T/T, TCF4_A/A, TCF4_A/T, TCF4_T/T, CDH20_A/A, CDH20_A/G, CDH20_G/G", size=10)
    add_body(doc, (
        "El sistema genera predicciones para cada fila, incluyendo probabilidad de riesgo, clasificación "
        "y valores SHAP individuales. Los resultados se pueden descargar como CSV."
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. MODELOS DE MACHINE LEARNING
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Modelos de Machine Learning", level=1)
    add_body(doc, (
        "ANXRISK utiliza dos modelos de aprendizaje automático basados en XGBoost (Gradient Boosting): "
        "un modelo estándar con 13 features clínicas y un modelo extendido con 22 features que incluye datos genéticos. "
        "Los modelos están serializados "
        "en formato .joblib y se cargan en tiempo de ejecución."
    ))

    add_heading_styled(doc, "7.1 Modelo Estándar — XGBoost (13 features)", level=2)
    add_body(doc, "Archivo: src/models/anxrisk_best_standard.joblib", bold=True)
    add_body(doc, "Se utiliza cuando el profesional NO incluye datos genéticos en la evaluación. Las 13 variables son:")
    add_table(doc,
        ["#", "Variable", "Descripción", "Valores"],
        [
            ["1", "EDAD24", "Grupo de edad", "0 (≤24 años) / 1 (>24 años)"],
            ["2", "AEFGROUPS", "Educación formal", "0 (≤14 años) / 1 (≥15 años)"],
            ["3", "LTE12_0", "Sin eventos estresantes", "0 / 1"],
            ["4", "LTE12_1", "Un evento estresante", "0 / 1"],
            ["5", "LTE12_2", "Dos o más eventos", "0 / 1"],
            ["6", "SF12F_Q1", "Salud física cuartil 1", "0 / 1"],
            ["7", "SF12F_Q2", "Salud física cuartil 2", "0 / 1"],
            ["8", "SF12F_Q3", "Salud física cuartil 3", "0 / 1"],
            ["9", "SF12F_Q4", "Salud física cuartil 4", "0 / 1"],
            ["10", "SF12M_Q1", "Salud mental cuartil 1", "0 / 1"],
            ["11", "SF12M_Q2", "Salud mental cuartil 2", "0 / 1"],
            ["12", "SF12M_Q3", "Salud mental cuartil 3", "0 / 1"],
            ["13", "SF12M_Q4", "Salud mental cuartil 4", "0 / 1"],
        ],
        col_widths=[1, 3, 5, 5]
    )

    add_heading_styled(doc, "7.2 Modelo Extendido — XGBoost (22 features)", level=2)
    add_body(doc, "Archivo: src/models/anxrisk_best_extended.joblib", bold=True)
    add_body(doc, "Se utiliza cuando el profesional SÍ incluye datos genéticos. Incluye las 13 features estándar más 9 variables genéticas (one-hot encoding de los 3 genes):")
    add_table(doc,
        ["#", "Variable", "Gen", "Genotipo"],
        [
            ["14", "PRKCA_C/C", "PRKCA", "C/C"],
            ["15", "PRKCA_C/T", "PRKCA", "C/T"],
            ["16", "PRKCA_T/T", "PRKCA", "T/T"],
            ["17", "TCF4_A/A", "TCF4", "A/A"],
            ["18", "TCF4_A/T", "TCF4", "A/T"],
            ["19", "TCF4_T/T", "TCF4", "T/T"],
            ["20", "CDH20_A/A", "CDH20", "A/A"],
            ["21", "CDH20_A/G", "CDH20", "A/G"],
            ["22", "CDH20_G/G", "CDH20", "G/G"],
        ],
        col_widths=[1, 4, 3, 3]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. INTERPRETABILIDAD SHAP
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. Interpretabilidad SHAP", level=1)
    add_body(doc, (
        "ANXRISK utiliza SHAP (SHapley Additive exPlanations) para proporcionar interpretabilidad "
        "individual de cada predicción. Se emplea TreeExplainer para ambos modelos XGBoost "
        "(estándar y extendido)."
    ))
    add_body(doc, "¿Cómo leer el gráfico SHAP?", bold=True)
    add_bullet(doc, "Factores que AUMENTAN el riesgo de ansiedad del paciente.", bold_prefix="Barras rojas (derecha): ")
    add_bullet(doc, "Factores que DISMINUYEN el riesgo de ansiedad del paciente.", bold_prefix="Barras azules (izquierda): ")
    add_bullet(doc, "Cuanto más larga la barra, mayor es el impacto de ese factor en la predicción individual.", bold_prefix="Tamaño de la barra: ")
    add_bullet(doc, "El punto de partida es el valor base (E[f(x)]), que representa la predicción promedio del modelo.", bold_prefix="Valor base: ")
    add_bullet(doc, "La suma del valor base más todos los valores SHAP da como resultado la predicción final (f(x)).", bold_prefix="Predicción final: ")

    add_body(doc, (
        "El gráfico SHAP permite al profesional entender QUÉ factores están contribuyendo más "
        "al nivel de riesgo de cada paciente específico, facilitando intervenciones personalizadas."
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. CLASIFICACIÓN DEL NIVEL DE RIESGO
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. Clasificación del Nivel de Riesgo", level=1)
    add_body(doc, (
        "La probabilidad predicha por el modelo se clasifica en tres niveles de riesgo "
        "mediante umbrales fijos:"
    ))
    add_table(doc,
        ["Nivel de Riesgo", "Rango de Probabilidad", "Color", "Recomendación Clínica"],
        [
            ["Bajo", "0.00 – 0.29", "🟢 Verde", "Sin intervención inmediata requerida. Seguimiento de rutina."],
            ["Moderado", "0.30 – 0.69", "🟡 Amarillo", "Monitoreo activo. Evaluación de seguimiento recomendada."],
            ["Alto", "0.70 – 1.00", "🔴 Rojo", "Evaluación profesional prioritaria. Intervención recomendada."],
        ],
        col_widths=[3, 3.5, 3, 6.5]
    )
    add_body(doc, (
        "Estos umbrales son fijos y no dependen de un cálculo de Youden dinámico. "
        "Fueron establecidos para proporcionar claridad clínica en la toma de decisiones."
    ), italic=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. GENERACIÓN DEL REPORTE PDF
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. Generación del Reporte PDF", level=1)
    add_body(doc, (
        "El reporte PDF se genera con ReportLab y contiene toda la información de la evaluación "
        "en un formato profesional listo para su archivo clínico. El documento incluye:"
    ))
    add_body(doc, "Secciones del PDF:", bold=True)
    add_bullet(doc, "Logo ANXRISK, nombre del paciente, documento de identidad, fecha de evaluación.", bold_prefix="1. Portada: ")
    add_bullet(doc, "Nombre, documento, edad, género, años de educación.", bold_prefix="2. Datos demográficos: ")
    add_bullet(doc, "Puntaje, clasificación y detalle de cada evento.", bold_prefix="3. Eventos vitales (LTE-12): ")
    add_bullet(doc, "Puntajes PCS y MCS con cuartiles.", bold_prefix="4. Calidad de vida (SF-12): ")
    add_bullet(doc, "Puntaje HADS y clasificación.", bold_prefix="5. Ansiedad HADS: ")
    add_bullet(doc, "Puntaje normalizado ZSAS y clasificación.", bold_prefix="6. Ansiedad ZSAS: ")
    add_bullet(doc, "Genotipos seleccionados (si aplica).", bold_prefix="7. Panel genético: ")
    add_bullet(doc, "Probabilidad, nivel de riesgo, modelo utilizado.", bold_prefix="8. Resultado de la predicción: ")
    add_bullet(doc, "Gráfico waterfall SHAP embebido como imagen.", bold_prefix="9. Análisis SHAP: ")
    add_bullet(doc, "Espacio para firma, nombre del profesional, cargo, institución, tarjeta profesional.", bold_prefix="10. Firma del profesional: ")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 11. DEPENDENCIAS Y VERSIONES
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "11. Dependencias y Versiones", level=1)
    add_body(doc, (
        "Las versiones de las dependencias están fijadas para garantizar compatibilidad total "
        "con los archivos .joblib de los modelos entrenados."
    ))
    add_table(doc,
        ["Librería", "Versión", "Propósito"],
        [
            ["Python", "3.12.12", "Lenguaje de programación (runtime.txt)"],
            ["xgboost", "≥2.0.0", "Modelo de clasificación XGBoost (estándar y extendido)"],
            ["scikit-learn", "1.6.1", "Preprocesamiento y métricas de ML"],
            ["shap", "0.51.0", "Interpretabilidad individual (TreeExplainer)"],
            ["joblib", "1.5.3", "Serialización/deserialización de modelos"],
            ["numpy", "2.0.2", "Operaciones numéricas"],
            ["pandas", "2.2.2", "Manipulación de DataFrames"],
            ["scipy", "1.16.3", "Funciones científicas"],
            ["streamlit", "≥1.28.0", "Framework web de la aplicación"],
            ["reportlab", "≥4.0.0", "Generación de reportes PDF"],
            ["matplotlib", "≥3.7.0", "Gráficos SHAP"],
            ["openpyxl", "≥3.1.0", "Soporte para archivos Excel"],
        ],
        col_widths=[3.5, 3, 9]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 12. PRIVACIDAD Y PROTECCIÓN DE DATOS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "12. Privacidad y Protección de Datos", level=1)
    add_body(doc, (
        "ANXRISK cumple con la normatividad colombiana de protección de datos personales:"
    ))
    add_bullet(doc, "Régimen General de Protección de Datos Personales (Colombia)", bold_prefix="Ley 1581 de 2012: ")
    add_bullet(doc, "Reglamentario de la Ley 1581 de 2012 (Habeas Data)", bold_prefix="Decreto 1377 de 2013: ")
    add_bullet(doc, "Habeas Data", bold_prefix="Ley 1266 de 2008: ")
    add_bullet(doc, "Investigación en salud", bold_prefix="Resolución 8430 de 1993: ")

    add_body(doc, "Principios de tratamiento de datos:", bold=True)
    add_bullet(doc, "Los datos se procesan exclusivamente en la sesión del navegador del usuario.", bold_prefix="Procesamiento local: ")
    add_bullet(doc, "No se almacenan datos personales en servidores externos ni bases de datos permanentes.", bold_prefix="Sin almacenamiento externo: ")
    add_bullet(doc, "Al cerrar la sesión del navegador, todos los datos se eliminan automáticamente.", bold_prefix="Eliminación automática: ")
    add_bullet(doc, "Los reportes PDF descargados quedan bajo la custodia del profesional responsable.", bold_prefix="Custodia del profesional: ")
    add_bullet(doc, "El profesional de salud debe contar con la autorización previa, expresa e informada del paciente.", bold_prefix="Consentimiento informado: ")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 13. PREGUNTAS FRECUENTES
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "13. Preguntas Frecuentes", level=1)

    faqs = [
        (
            "¿Puedo usar ANXRISK sin datos genéticos?",
            "Sí. Los datos genéticos son un módulo de confirmación opcional. Si no se incluyen, "
            "el sistema utiliza el modelo estándar de 13 features. Los resultados siguen siendo "
            "clínicamente válidos."
        ),
        (
            "¿Los datos del paciente se guardan en algún servidor?",
            "No. Todos los datos se procesan exclusivamente en la sesión del navegador. Al cerrar "
            "la sesión o la pestaña, todos los datos se eliminan automáticamente. Solo el reporte PDF "
            "descargado persiste, bajo custodia del profesional."
        ),
        (
            "¿Qué significan las barras del gráfico SHAP?",
            "Las barras rojas indican factores que aumentan el riesgo y las azules factores que lo disminuyen. "
            "El tamaño de cada barra indica la magnitud del impacto de ese factor en la predicción individual."
        ),
        (
            "¿Puedo evaluar múltiples pacientes a la vez?",
            "Sí, mediante la funcionalidad de Análisis Masivo. Cargue un archivo CSV con las columnas "
            "requeridas y el sistema generará predicciones para cada fila."
        ),
        (
            "¿Qué modelo de machine learning utiliza ANXRISK?",
            "ANXRISK utiliza dos modelos basados en XGBoost (Gradient Boosting): uno estándar (13 features clínicas) "
            "y uno extendido con datos genéticos (22 features). "
            "Los modelos fueron entrenados y validados en Google Colab."
        ),
        (
            "¿Cómo se determina el nivel de riesgo?",
            "La probabilidad predicha por el modelo se clasifica en tres niveles con umbrales fijos: "
            "Bajo (< 0.30), Moderado (0.30 – 0.69), Alto (≥ 0.70)."
        ),
        (
            "¿La edad del paciente afecta la predicción?",
            "Sí. La edad se transforma en una variable binaria: ≤24 años (grupo 0) y >24 años (grupo 1). "
            "Este punto de corte fue establecido durante el entrenamiento del modelo."
        ),
    ]

    for pregunta, respuesta in faqs:
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"P: {pregunta}")
        run_q.bold = True
        run_q.font.size = Pt(11)
        run_q.font.color.rgb = AZUL_PRIMARIO

        p_a = doc.add_paragraph()
        run_a = p_a.add_run(f"R: {respuesta}")
        run_a.font.size = Pt(11)
        run_a.font.color.rgb = GRIS_OSCURO

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 14. HISTORIAL DE CAMBIOS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "14. Historial de Cambios (Versión 5)", level=1)
    add_body(doc, "Cambios principales respecto a la versión anterior:", bold=True)

    add_table(doc,
        ["Categoría", "Cambio", "Detalle"],
        [
            ["Nuevo campo", "Documento de identidad", "Se agregó el campo 'Documento de identidad' al formulario demográfico, al resumen de resultados, a la portada del PDF y a la tabla de datos demográficos del PDF."],
            ["Limpieza de código", "Eliminación de modelos inexistentes", "Se eliminaron todas las referencias a modelos obsoletos. Se usa XGBoost para ambos modelos: estándar (13 features) y extendido (22 features)."],
            ["Limpieza de código", "Funciones no usadas eliminadas", "Se eliminaron las funciones transformar_genotipo_prkca, transformar_genotipo_tcf4, transformar_genotipo_cdh20, youden_threshold y validar_años_educacion de calculos.py."],
            ["Limpieza de código", "Imports no usados eliminados", "Se eliminó 'import os' de resultados.py y 'obtener_registro_actual' de datos_geneticos.py. Se eliminó código comentado obsoleto."],
            ["Limpieza de código", "Archivos obsoletos eliminados", "Se eliminaron main.css.bak, main.css.old y directorios __pycache__."],
            ["Documentación", "Docstrings mejorados", "Se agregó docstring de módulo a calculos.py describiendo los 5 bloques funcionales."],
            ["Dependencias", "Versiones fijadas", "Se fijaron versiones en requirements.txt: xgboost>=2.0.0, scikit-learn==1.6.1, shap==0.51.0, numpy==2.0.2, etc."],
            ["Runtime", "Python actualizado", "Se actualizó runtime.txt de python-3.11.0 a python-3.12.12, alineándolo con la versión de Colab."],
            ["SHAP", "Explainer unificado", "Se usa TreeExplainer para ambos modelos XGBoost (estándar y extendido)."],
        ],
        col_widths=[3, 4, 9]
    )

    # ── Guardar ──
    output_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "docs",
        "ANXRISK_Manual_de_Usuario_v5.docx"
    )
    doc.save(output_path)
    print(f"✅ Manual generado: {output_path}")
    return output_path


if __name__ == "__main__":
    generar_manual()
