#!/usr/bin/env python3
"""
Genera el Manual de Usuario ANXRISK en formato DOCX profesional
a partir del archivo Markdown docs/USER_MANUAL.md
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Colores del diseño ANXRISK ──────────────────────────────────
AMBER = RGBColor(0xD4, 0x91, 0x1D)       # #D4911D
DARK_TEXT = RGBColor(0x2D, 0x2D, 0x2D)    # #2D2D2D
SURFACE = RGBColor(0xF0, 0xED, 0xEA)      # #F0EDEA
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, left, right with dict {sz, color, val}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return run


def create_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        set_cell_shading(cell, "D4911D")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_TEXT
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F9F7F5")
            else:
                set_cell_shading(cell, "FFFFFF")

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def parse_markdown_table(lines):
    """Parse a markdown table from lines, returning headers and rows."""
    headers = []
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(set(c) <= {'-', ':', ' '} for c in cells):
            continue  # separator line
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def process_inline_formatting(paragraph, text, base_size=10):
    """Process inline markdown formatting (**bold**, *italic*, `code`)."""
    # Pattern to match bold, italic, code, or links
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.size = Pt(base_size)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_TEXT

        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.size = Pt(base_size)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_TEXT
        elif match.group(3):  # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True
            run.font.size = Pt(base_size)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_TEXT
        elif match.group(4):  # Code
            run = paragraph.add_run(match.group(4))
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
        elif match.group(5):  # Link
            run = paragraph.add_run(match.group(5))
            run.font.size = Pt(base_size)
            run.font.name = 'Calibri'
            run.font.color.rgb = AMBER
            run.underline = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(base_size)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_TEXT


def add_blockquote(doc, text):
    """Add a styled blockquote paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Add left border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="8" w:color="D4911D"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    # Clean the text
    clean_text = text.lstrip('> ').strip()
    process_inline_formatting(p, clean_text, base_size=9)
    for run in p.runs:
        run.italic = True
        run.font.color.rgb = MID_GRAY


def add_code_block(doc, code_lines):
    """Add a styled code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    # Background shading
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shd)
    
    code_text = '\n'.join(code_lines)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK_TEXT


def generate_docx(md_path, output_path):
    """Main function: parse Markdown and generate DOCX."""
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # ── Page setup ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # ── Default style ──────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK_TEXT
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15
    
    # ── Heading styles ─────────────────────────────────────────
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = AMBER if level <= 2 else DARK_TEXT
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
        heading_style.paragraph_format.space_after = Pt(6)
        
        if level == 1:
            heading_style.font.size = Pt(22)
        elif level == 2:
            heading_style.font.size = Pt(16)
        elif level == 3:
            heading_style.font.size = Pt(13)
        else:
            heading_style.font.size = Pt(11)
    
    # ── Cover page ─────────────────────────────────────────────
    # Spacer
    for _ in range(4):
        doc.add_paragraph()
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('ANXRISK')
    run.bold = True
    run.font.size = Pt(42)
    run.font.color.rgb = AMBER
    run.font.name = 'Calibri'
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('Sistema de Estratificación del Riesgo\nde Trastornos de Ansiedad')
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_TEXT
    run.font.name = 'Calibri'
    
    # Divider line
    div_p = doc.add_paragraph()
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div_p.add_run('━' * 40)
    run.font.color.rgb = AMBER
    run.font.size = Pt(14)
    
    # Manual title
    manual_p = doc.add_paragraph()
    manual_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    manual_p.paragraph_format.space_before = Pt(20)
    run = manual_p.add_run('Manual de Usuario Completo')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = DARK_TEXT
    run.font.name = 'Calibri'
    
    # Version info
    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.paragraph_format.space_before = Pt(30)
    run = ver_p.add_run('Versión 2.0 — 2025')
    run.font.size = Pt(11)
    run.font.color.rgb = MID_GRAY
    run.font.name = 'Calibri'
    
    # Author
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = auth_p.add_run('Breyner Joel Quiñones Castro')
    run.font.size = Pt(11)
    run.font.color.rgb = MID_GRAY
    run.font.name = 'Calibri'
    
    # Page break after cover
    doc.add_page_break()
    
    # ── Parse content ──────────────────────────────────────────
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    skip_main_title = True  # Skip the first H1 since we have cover page
    
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip('\n')
        stripped = raw.strip()
        
        # ── Code blocks ────────────────────────────────────────
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table:
                    headers, rows = parse_markdown_table(table_lines)
                    if headers:
                        create_styled_table(doc, headers, rows)
                    table_lines = []
                    in_table = False
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue
        
        # ── Tables ─────────────────────────────────────────────
        if stripped.startswith('|') and '|' in stripped[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            headers, rows = parse_markdown_table(table_lines)
            if headers:
                create_styled_table(doc, headers, rows)
            table_lines = []
            in_table = False
        
        # ── Empty lines ───────────────────────────────────────
        if not stripped:
            i += 1
            continue
        
        # ── Horizontal rules ──────────────────────────────────
        if stripped in ('---', '***', '___'):
            i += 1
            continue
        
        # ── Headings ──────────────────────────────────────────
        heading_match = re.match(r'^(#{1,5})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            
            # Skip the main title (first H1) since we have cover page
            if level == 1 and skip_main_title:
                skip_main_title = False
                i += 1
                continue
            
            # Clean text of markdown formatting for heading
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            clean = re.sub(r'`(.+?)`', r'\1', clean)
            clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean)
            
            heading_level = min(level, 4)
            heading = doc.add_heading(clean, level=heading_level)
            
            # Add amber underline for H2
            if level == 2:
                border_p = doc.add_paragraph()
                border_p.paragraph_format.space_before = Pt(0)
                border_p.paragraph_format.space_after = Pt(8)
                run = border_p.add_run('━' * 60)
                run.font.size = Pt(6)
                run.font.color.rgb = AMBER
            
            i += 1
            continue
        
        # ── Blockquotes ───────────────────────────────────────
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('> ').strip()
            # Collect multi-line blockquotes
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith('>'):
                quote_text += ' ' + lines[j].strip().lstrip('> ').strip()
                j += 1
            add_blockquote(doc, quote_text)
            i = j
            continue
        
        # ── Bullet lists ──────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', stripped)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            
            p = doc.add_paragraph()
            p.style = doc.styles['List Bullet']
            p.paragraph_format.left_indent = Cm(1.0 + (indent // 2) * 0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            # Clear default text
            p.clear()
            process_inline_formatting(p, text, base_size=10)
            
            i += 1
            continue
        
        # ── Numbered lists ────────────────────────────────────
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph()
            p.style = doc.styles['List Number']
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.clear()
            process_inline_formatting(p, text, base_size=10)
            i += 1
            continue
        
        # ── Regular paragraphs ────────────────────────────────
        p = doc.add_paragraph()
        process_inline_formatting(p, stripped, base_size=10)
        i += 1
    
    # Flush pending table
    if in_table:
        headers, rows = parse_markdown_table(table_lines)
        if headers:
            create_styled_table(doc, headers, rows)
    
    # ── Footer ─────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run('© 2025 Breyner Joel Quiñones Castro. Todos los derechos reservados.')
    run.font.size = Pt(8)
    run.font.color.rgb = MID_GRAY
    run.font.name = 'Calibri'
    run.italic = True
    
    # ── Save ───────────────────────────────────────────────────
    doc.save(output_path)
    print(f"✅ Documento generado exitosamente: {output_path}")
    print(f"   Tamaño: {os.path.getsize(output_path) / 1024:.1f} KB")


if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(base_dir, 'docs', 'USER_MANUAL.md')
    output_path = os.path.join(base_dir, 'docs', 'ANXRISK_Manual_de_Usuario.docx')
    
    generate_docx(md_path, output_path)
